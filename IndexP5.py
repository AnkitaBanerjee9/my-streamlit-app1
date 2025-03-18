import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO

# Configure the page
st.set_page_config(page_title="Performance Report Comparison", layout="wide")
st.title("Performance Report Analysis")

# Sidebar: File upload
st.sidebar.header("Upload Test Report")
uploaded_file = st.sidebar.file_uploader("Upload the report file", type=["csv", "xlsx"])

@st.cache_data
def load_data(file):
    if file is not None:
        if file.name.endswith('.csv'):
            return pd.read_csv(file)
        else:
            return pd.read_excel(file)
    return None

df = load_data(uploaded_file)

if df is not None:
    # Section 1: Report Preview
    st.header("Report Preview")
    st.dataframe(df)
    
    # Section 2: Filtering Table
    st.header("Filtered Data Table")
    filter_columns = st.sidebar.multiselect("Select columns for filtering", df.columns.tolist(), default=df.columns.tolist())
    logical_column = st.sidebar.selectbox("Select column for row filtering", df.columns.tolist())
    unique_vals = df[logical_column].dropna().unique().tolist()
    selected_vals = st.sidebar.multiselect("Select row values to display", unique_vals, default=unique_vals)
    
    if logical_column in df.columns:
        filtered_df = df[df[logical_column].isin(selected_vals)][filter_columns]
    else:
        filtered_df = df[filter_columns]
    
    st.dataframe(filtered_df)
    
    # Section 3: Download Filtered Data
    st.sidebar.subheader("Download Filtered Data")
    file_format = st.sidebar.radio("Select format", ["CSV", "Excel"])
    
    if st.sidebar.button("Download"):
        if not filtered_df.empty:
            if file_format == "CSV":
                filtered_df.to_csv("filtered_data.csv", index=False)
                st.sidebar.success("Filtered data saved as filtered_data.csv")
            else:
                filtered_df.to_excel("filtered_data.xlsx", index=False)
                st.sidebar.success("Filtered data saved as filtered_data.xlsx")
        else:
            st.sidebar.error("No data available to download.")
    
    if st.sidebar.button("View Downloaded Report"):
        st.write("Displaying the downloaded report:")
        st.dataframe(filtered_df)
    
    # Section 4: Response Time Comparison
    st.header("Response Time Comparison: Run1 vs Run2 vs Run3")
    required_cols = ['TransactionName', 'SLA', 'Run1-90Percent', 'Run2-90Percent', 'Run3-90Percent']
    available_cols = [col for col in required_cols if col in filtered_df.columns]
    
    if 'TransactionName' in available_cols and len(available_cols) > 1:
        avg_response_times = {col: filtered_df[col].mean() for col in available_cols if 'Run' in col}
        
        for run, avg in avg_response_times.items():
            st.write(f"**{run} Average Response Time:** {avg:.2f}")
        
        if avg_response_times:
            best_run = min(avg_response_times, key=avg_response_times.get)
            reason = f"{best_run} is the best because it has the lowest average response time of {avg_response_times[best_run]:.2f}."
            st.success(f"{best_run} has the best (lowest) response times overall. {reason}")
        else:
            st.warning("No valid data for response time comparison.")
    else:
        st.warning("Not enough data for response time comparison.")
    
    # Section 5: Graphical Comparison
    if 'TransactionName' in filtered_df.columns:
        st.subheader("Graphical Comparison by Transaction")
        transaction_options = filtered_df['TransactionName'].dropna().unique().tolist()
        selected_transactions = st.multiselect("Select transactions to display in graph", transaction_options, default=transaction_options if transaction_options else [])
        
        if selected_transactions:
            df_graph = filtered_df[filtered_df['TransactionName'].isin(selected_transactions)]
            df_plot = df_graph.melt(id_vars='TransactionName', value_vars=[col for col in available_cols if 'Run' in col], var_name='Run', value_name='Response Time')
            
            if not df_plot.empty:
                fig = px.bar(df_plot, x='TransactionName', y='Response Time', color='Run', barmode='group', title="Response Time Comparison per Transaction")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.warning("No data available for graphing.")
    
    # Section 6: Anomaly Detection
    if 'SLA' in filtered_df.columns and any(col in filtered_df.columns for col in ['Run1-90Percent', 'Run2-90Percent', 'Run3-90Percent']):
        st.subheader("Anomaly Detection: Transactions Exceeding SLA")
        slow_transactions = filtered_df[filtered_df.filter(like='Run').gt(filtered_df['SLA'], axis=0).any(axis=1)]
        
        if not slow_transactions.empty:
            st.warning("The following transactions exceed the SLA:")
            st.dataframe(slow_transactions[['TransactionName', 'SLA'] + [col for col in available_cols if 'Run' in col]])
        else:
            st.success("No transactions exceed SLA limits.")
    
    # Section 7: Performance Trend Analysis
    if 'TransactionName' in filtered_df.columns:
        st.subheader("Performance Trend Analysis")
        df_trend = filtered_df.melt(id_vars='TransactionName', value_vars=[col for col in available_cols if 'Run' in col], var_name='Run', value_name='Response Time')
        
        if not df_trend.empty:
            fig_trend = px.line(df_trend, x='TransactionName', y='Response Time', color='Run', markers=True, title="Response Time Trend Over Runs")
            st.plotly_chart(fig_trend, use_container_width=True)
        else:
            st.warning("No data available for trend analysis.")
    
    # Section 8: Generate Word Report
    if st.sidebar.button("Generate Word Report"):
        doc = Document()
        doc.add_heading("Performance Report", level=1)
        
        doc.add_heading("Filtered Data Table", level=2)
        if not filtered_df.empty:
            table = doc.add_table(rows=1, cols=len(filtered_df.columns))
            hdr_cells = table.rows[0].cells
            for j, col in enumerate(filtered_df.columns):
                hdr_cells[j].text = col
            
            for i, row in filtered_df.iterrows():
                row_cells = table.add_row().cells
                for j, value in enumerate(row):
                    row_cells[j].text = str(value)
            
            doc.save("Performance_Report.docx")
            st.sidebar.success("Word document generated: Performance_Report.docx")
        else:
            st.sidebar.error("No data available to generate report.")
else:
    st.info("Please upload a report file to begin the analysis.")
